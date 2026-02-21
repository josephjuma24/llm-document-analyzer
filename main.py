"""
LLM Document Analyzer — main.py
A Streamlit app that parses uploaded documents and runs LLM-powered analysis.
"""

import io
import os
import traceback
from typing import Optional

import streamlit as st

# ─── Page Config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="LLM Document Analyzer",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Helpers: Document Parsing ───────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber (with PyPDF2 as fallback)."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # also extract tables
                for table in page.extract_tables():
                    rows = ["\t".join(str(c) if c else "" for c in row) for row in table]
                    text += "\n" + "\n".join(rows)
                pages.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(pages)
    except Exception:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n\n".join(
                f"[Page {i+1}]\n{page.extract_text() or ''}"
                for i, page in enumerate(reader.pages)
            )
        except Exception as e:
            return f"[PDF parsing error: {e}]"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX parsing error: {e}]"


def parse_excel(file_bytes: bytes, suffix: str) -> str:
    """Extract text from Excel / CSV files via pandas."""
    try:
        import pandas as pd
        if suffix == ".csv":
            df_dict = {"Sheet1": pd.read_csv(io.BytesIO(file_bytes))}
        else:
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        parts = []
        for sheet_name, df in df_dict.items():
            parts.append(f"[Sheet: {sheet_name}]\n{df.to_string(index=False)}")
        return "\n\n".join(parts)
    except Exception as e:
        return f"[Excel/CSV parsing error: {e}]"


def parse_text(file_bytes: bytes) -> str:
    """Fallback: decode as UTF-8 text."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Text parsing error: {e}]"


def extract_content(uploaded_file) -> tuple[str, str]:
    """
    Given a Streamlit UploadedFile, return (content, preview_content).
    content  — full extracted text (possibly truncated for LLM)
    preview  — first ~2000 chars for display
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        content = parse_pdf(file_bytes)
    elif name.endswith((".docx", ".doc")):
        content = parse_docx(file_bytes)
    elif name.endswith((".xlsx", ".xls", ".csv")):
        suffix = os.path.splitext(name)[1]
        content = parse_excel(file_bytes, suffix)
    else:
        content = parse_text(file_bytes)

    preview = content[:2000] + ("…" if len(content) > 2000 else "")
    return content, preview


# ─── Helpers: LLM Integration ────────────────────────────────────────────────

MAX_CONTENT_CHARS = 12_000  # ~3 k tokens; keeps costs low


def get_api_key() -> Optional[str]:
    """Read API key from Streamlit secrets or environment variable."""
    try:
        return st.secrets["OPENAI_API_KEY"]
    except Exception:
        return os.environ.get("OPENAI_API_KEY", "")


def call_llm(system_prompt: str, user_prompt: str, api_key: str) -> str:
    """Call OpenAI Chat Completions API and return the assistant reply."""
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
        max_tokens=1500,
    )
    return response.choices[0].message.content


def build_prompt(task: str, content: str, query: str = "") -> tuple[str, str]:
    """Return (system_prompt, user_prompt) for the chosen task."""
    truncated = content[:MAX_CONTENT_CHARS]
    sys = "You are an expert document analyst. Answer clearly and concisely using only the document content."

    if task == "Summarize":
        user = f"Summarize the following document in 3–5 paragraphs:\n\n{truncated}"
    elif task == "Extract Entities":
        user = (
            f"Extract key entities from the document: people, organizations, "
            f"dates, locations, and monetary values. Present as a structured list.\n\n{truncated}"
        )
    elif task == "Q&A":
        user = f"Answer the following question based solely on the document:\n\nQuestion: {query}\n\nDocument:\n{truncated}"
    elif task == "Sentiment Analysis":
        user = f"Analyze the overall sentiment (positive/negative/neutral) and tone of the document. Explain your reasoning.\n\n{truncated}"
    elif task == "Compare Documents":
        user = f"Compare and contrast these documents, highlighting similarities and key differences:\n\n{truncated}"
    elif task == "Custom Query":
        user = f"Based on the document below, {query}\n\nDocument:\n{truncated}"
    else:
        user = truncated

    return sys, user


def analyze(task: str, combined_content: str, query: str, api_key: str) -> str:
    """Run the LLM analysis and return the result string."""
    sys_prompt, user_prompt = build_prompt(task, combined_content, query)
    return call_llm(sys_prompt, user_prompt, api_key)


# ─── UI ──────────────────────────────────────────────────────────────────────

def main():
    # Sidebar
    with st.sidebar:
        st.title("⚙️ Settings")
        api_key = get_api_key()
        if not api_key:
            api_key = st.text_input("OpenAI API Key", type="password",
                                     help="Enter your OpenAI key or set it via Streamlit Secrets.")
        else:
            st.success("API key loaded from secrets ✓")

        st.markdown("---")
        st.markdown("**Supported formats**")
        st.markdown("PDF · DOCX · XLSX · XLS · CSV · TXT · MD")
        st.markdown("---")
        st.markdown("Built with [Streamlit](https://streamlit.io) + OpenAI")

    # Main area
    st.title("🔍 LLM Document Analyzer")
    st.markdown(
        "Upload one or more documents, choose an analysis type, and let the LLM do the heavy lifting."
    )

    # File upload
    uploaded_files = st.file_uploader(
        "Upload document(s)",
        type=["pdf", "docx", "doc", "xlsx", "xls", "csv", "txt", "md"],
        accept_multiple_files=True,
    )

    if not uploaded_files:
        st.info("👆 Upload at least one file to get started.")
        return

    # Parse files
    documents: dict[str, tuple[str, str]] = {}  # name → (full_content, preview)
    for uf in uploaded_files:
        if uf.size > 20 * 1024 * 1024:  # 20 MB guard
            st.warning(f"⚠️ **{uf.name}** is larger than 20 MB and will be skipped.")
            continue
        with st.spinner(f"Parsing {uf.name}…"):
            try:
                content, preview = extract_content(uf)
                documents[uf.name] = (content, preview)
            except Exception:
                st.error(f"Failed to parse **{uf.name}**:\n```\n{traceback.format_exc()}\n```")

    if not documents:
        return

    # Document preview tabs
    st.markdown("### 📄 Document Previews")
    tabs = st.tabs(list(documents.keys()))
    for tab, (name, (_, preview)) in zip(tabs, documents.items()):
        with tab:
            st.text_area("Content preview (first 2 000 chars)", preview,
                         height=200, disabled=True, key=f"preview_{name}")

    st.markdown("---")

    # Analysis options
    st.markdown("### 🧠 Analysis")
    col1, col2 = st.columns([1, 2])

    with col1:
        task_options = ["Summarize", "Extract Entities", "Q&A",
                        "Sentiment Analysis", "Compare Documents", "Custom Query"]
        task = st.selectbox("Choose analysis type", task_options)

    query = ""
    with col2:
        if task in ("Q&A", "Custom Query"):
            query = st.text_area(
                "Your question / query",
                placeholder="Ask anything about the document(s)…",
                height=100,
            )

    run = st.button("▶ Run Analysis", type="primary", use_container_width=True)

    if run:
        if not api_key:
            st.error("Please provide an OpenAI API key in the sidebar.")
            return
        if task in ("Q&A", "Custom Query") and not query.strip():
            st.warning("Please enter a question or query before running.")
            return

        # Combine all document content
        if len(documents) == 1:
            combined = list(documents.values())[0][0]
        else:
            combined = "\n\n".join(
                f"=== Document: {name} ===\n{content}"
                for name, (content, _) in documents.items()
            )

        with st.spinner("Calling LLM… this may take a few seconds."):
            try:
                result = analyze(task, combined, query, api_key)
                st.markdown("### 📊 Result")
                st.markdown(result)
            except Exception as e:
                err = str(e)
                if "api_key" in err.lower() or "authentication" in err.lower():
                    st.error("❌ Invalid or missing API key. Please check your key and try again.")
                elif "rate limit" in err.lower():
                    st.error("❌ Rate limit reached. Wait a moment and retry.")
                elif "context_length" in err.lower() or "maximum context" in err.lower():
                    st.error("❌ Document is too long even after truncation. Try a shorter document.")
                else:
                    st.error(f"❌ LLM error: {e}")


if __name__ == "__main__":
    main()
